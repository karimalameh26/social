"""
Extract endpoint information from the documentation files
"""
import docx
import re
import json

def extract_endpoints_from_docx(file_path):
    """Extract endpoint URLs and information from a .docx file"""
    doc = docx.Document(file_path)
    endpoints = []
    
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    return full_text

def parse_endpoint_info(text_lines):
    """Parse endpoint information from text"""
    endpoints = {}
    
    for line in text_lines:
        # Look for URLs/endpoints
        urls = re.findall(r'https?://[^\s]+', line)
        if urls:
            endpoints[len(endpoints)] = {
                'url': urls[0],
                'description': line
            }
    
    return endpoints

if __name__ == "__main__":
    print("Extracting endpoints from documentation...")
    
    try:
        # Extract from both documents
        doc1_text = extract_endpoints_from_docx("X EndPoints 1.docx")
        doc2_text = extract_endpoints_from_docx("X EndPoints 2.docx")
        
        print("\n=== X EndPoints 1.docx ===")
        for line in doc1_text:
            print(line)
        
        print("\n=== X EndPoints 2.docx ===")
        for line in doc2_text:
            print(line)
        
        # Parse endpoints
        all_text = doc1_text + doc2_text
        endpoints = parse_endpoint_info(all_text)
        
        # Save to JSON
        with open('endpoints.json', 'w', encoding='utf-8') as f:
            json.dump(endpoints, f, indent=2, ensure_ascii=False)
        
        print(f"\nFound {len(endpoints)} endpoints")
        print("Saved to endpoints.json")
        
    except Exception as e:
        print(f"Error: {e}")
        print("\nPlease make sure python-docx is installed: pip install python-docx")

